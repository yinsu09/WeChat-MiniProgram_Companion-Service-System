# -*- coding: utf-8 -*-
"""按「报告案例：图书管理系统说明书」格式生成陪伴服务系统课设说明书"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, '学号_姓名_陪伴服务系统说明书.docx')

LINE_SPACING = 1.25
BODY_INDENT = Cm(0.74)


def set_run_font(run, latin_font, size, bold=False, east_asia=None):
    run.font.name = latin_font
    run.font.size = size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia or latin_font)


def apply_line_spacing(paragraph, spacing=LINE_SPACING):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = spacing


def add_body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    apply_line_spacing(p)
    if first_line_indent:
        p.paragraph_format.first_line_indent = BODY_INDENT
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', Pt(12), east_asia='宋体')
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    apply_line_spacing(p)
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', Pt(14), bold=True, east_asia='黑体')
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    apply_line_spacing(p)
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', Pt(12), bold=True, east_asia='黑体')
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    apply_line_spacing(p)
    p.paragraph_format.first_line_indent = BODY_INDENT
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', Pt(12), east_asia='宋体')
    return p


def add_figure_placeholder(doc, caption):
    add_body(doc, '【此处插入流程图/功能图/界面截图】', first_line_indent=False)
    p = doc.add_paragraph()
    apply_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_run_font(run, 'Times New Roman', Pt(10.5), east_asia='宋体')
    doc.add_paragraph()


def add_table_def(doc, title, desc, columns):
    add_body(doc, desc)
    p = doc.add_paragraph()
    apply_line_spacing(p)
    run = p.add_run(title)
    set_run_font(run, 'Times New Roman', Pt(12), bold=True, east_asia='宋体')

    headers = ['列名', '含义', '数据类型', '允许null值', '备注']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for r in para.runs:
                set_run_font(r, 'Times New Roman', Pt(10.5), bold=True, east_asia='宋体')

    for row in columns:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for para in cells[i].paragraphs:
                for r in para.runs:
                    set_run_font(r, 'Times New Roman', Pt(10.5), east_asia='宋体')
    doc.add_paragraph()


def add_numbered_list(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        apply_line_spacing(p)
        run = p.add_run(f'{i}. {item}')
        set_run_font(run, 'Times New Roman', Pt(12), east_asia='宋体')


def add_module_list(doc, items):
    for i, item in enumerate(items, 1):
        add_body(doc, f'{i}. {item}')


def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    for text in ['基于微信小程序的陪伴服务系统V1.0', '设计说明书']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_line_spacing(p)
        run = p.add_run(text)
        set_run_font(run, 'Times New Roman', Pt(22 if 'V1.0' in text else 18), bold=True, east_asia='黑体')
    for _ in range(6):
        doc.add_paragraph()
    for line in ['班级    ____________', '学号    ____________', '姓名    ____________']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_line_spacing(p)
        run = p.add_run(line)
        set_run_font(run, 'Times New Roman', Pt(14), east_asia='宋体')
    for _ in range(4):
        doc.add_paragraph()
    for line in ['天津城建大学', '计算机与信息工程学院', '数据库课程组', '二〇二六年六月二十日']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_line_spacing(p)
        run = p.add_run(line)
        set_run_font(run, 'Times New Roman', Pt(14), east_asia='宋体')
    doc.add_page_break()


def build_toc(doc):
    add_heading1(doc, '目  录（页码最后要更新）')
    toc = [
        ('第1章 系统分析', '1'),
        ('1.1 系统功能描述', '1'),
        ('1.2 系统功能模块设计', '2'),
        ('第2章 系统设计', '3'),
        ('2.1 系统整体业务流程图', '3'),
        ('2.2 系统各业务流程图', '4'),
        ('2.3 数据库设计', '8'),
        ('第3章 系统实现', '14'),
        ('3.1 系统开发环境', '14'),
        ('3.2 系统功能实现', '14'),
    ]
    for title, page in toc:
        p = doc.add_paragraph()
        apply_line_spacing(p)
        run = p.add_run(f'　　{title}\t{page}')
        set_run_font(run, 'Times New Roman', Pt(12), east_asia='宋体')
    doc.add_page_break()


def build_chapter1(doc):
    add_heading1(doc, '第1章 系统分析')
    add_heading2(doc, '1.1 系统功能描述')
    add_body(doc, '系统功能描述')
    add_body(doc, (
        '本微信小程序陪伴服务系统面向平台用户、服务人员以及系统管理员三类用户，依托微信小程序平台，'
        '实现服务展示、在线预约、订单管理、服务评价、积分营销以及后台运营管理等功能，为用户与服务人员之间'
        '建立安全、高效、便捷的服务平台。普通用户登录系统后，可浏览陪诊、陪护、陪玩、陪吃、陪游、陪学、陪聊等'
        '多种服务类型，并根据服务等级（铜/银/金）、价格及服务人员评分进行筛选，完成服务预约、在线下单、'
        '订单查询、退款申请、售后投诉以及服务评价等操作。系统支持用户自主选择服务人员，也支持根据服务人员等级、'
        '评价、工作模式及空闲状态进行智能指派，从而提高服务匹配效率。'
        '服务人员登录系统后，可查看平台分配或用户指定的订单，并完成接单、拒单、开始服务、暂停/恢复服务、'
        '结束服务等操作，同时还能够维护个人资料、发布与管理服务项目、设置可服务时间与休息时段、'
        '查看收入情况以及用户评价信息。管理员作为系统核心管理角色，可实现服务类型管理、服务项目管理、'
        '服务人员管理、订单监管与人工派单、退款审核、售后处理、营销活动设置（满减券、折扣券、组团游、新手礼包、'
        '限时折扣）以及数据统计分析等功能，并可根据服务人员完成订单数量与服务质量自动调整服务等级，'
        '对异常订单进行处理和管理。系统后端采用 Node.js + Express 框架，数据库采用 MySQL，'
        '实现了陪伴服务流程的信息化与智能化管理，有效解决了传统陪伴服务中信息分散、匹配效率低以及管理困难等问题。'
        '与此同时，系统提供热门服务统计、服务人员排行、用户消费行为分析、退款统计以及营收统计等数据分析功能，'
        '为平台运营优化提供数据支持，从而形成一个集服务展示、预约下单、评价反馈、智能匹配、积分营销以及后台管理'
        '于一体的综合性陪伴服务平台。'
    ))
    add_note(doc, '（用例图必须有，使用专用工具绘制）')
    add_figure_placeholder(doc, '用例图 陪伴服务系统')
    add_body(doc, '管理员功能需求：')
    add_numbered_list(doc, [
        '服务类型与服务项目管理：实现陪诊、陪护、陪玩、陪吃、陪游、陪学、陪聊等服务类型及具体服务项目的新增、修改、删除、上下架以及查询等功能，系统支持按服务类型分类管理，并可对服务价格、服务等级定价区间、服务时长及上下架状态进行设置。',
        '用户与服务人员信息管理：实现平台用户与服务人员的信息录入、修改、查询以及状态管理，对违规用户或服务人员可进行禁用与恢复操作，支持服务人员等级规则配置与自动升降级。',
        '订单与预约情况管理：实现订单信息的浏览、查询、人工派单与状态干预，对用户预约、服务进度、暂停/恢复以及退款申请进行统一监管。',
        '退款与售后处理：对用户退款申请进行审核（同意/拒绝），对售后申请（提前结束、纠纷、退款等）进行查看与处理，维护平台交易秩序。',
        '营销活动管理：管理满减券、折扣券、组团游活动、新手礼包、限时折扣等营销工具，支持活动创建、编辑、停用及效果统计。',
        '系统配置与派单规则管理：维护服务人员智能派单规则、退款规则、分级定价规则等系统配置参数。',
        '用户评价与反馈信息处理：对用户投诉、意见反馈、评价内容及售后信息进行查看与回复，维护平台服务质量。',
        '系统数据统计与分析功能：实现对平台订单数量、热门服务、服务人员排行、用户消费情况、退款率、营收数据以及等级分布等信息的统计与浏览，支持数据导出。',
    ])
    add_body(doc, '用户功能需求：')
    add_numbered_list(doc, [
        '服务信息的浏览与检索：用户可查询平台中的服务类型、服务项目、服务价格、服务等级以及服务人员信息，支持关键词搜索与条件筛选。',
        '个人信息的修改：用户可对个人资料、联系方式、服务区域偏好以及账号信息进行维护与更新，支持手机号绑定与密码修改。',
        '服务预约与下单功能：用户可在线选择服务类型与具体项目，填写预约时间、服务地点及定制需求，选择服务人员或由系统智能分配，完成订单提交与模拟支付。',
        '订单管理功能：用户可查看订单状态（待支付、待接单、待服务、服务中、已完成、已取消、退费中等），实现服务订单的取消、退款申请、确认完成、多次卡预约下次服务以及历史订单查询等操作。',
        '积分与优惠券功能：用户可通过消费获得积分，在积分商城兑换优惠券或礼品，查看积分明细与已领优惠券，下单时选择可用优惠券抵扣。',
        '服务评价与反馈功能：用户可对服务人员进行多维度评分与文字/图片评价，查看历史评价，并向平台提交售后申请或投诉信息。',
        '消息通知功能：用户可接收订单状态变更、系统公告、优惠活动等通知消息，并标记已读。',
    ])
    add_body(doc, '服务人员功能需求：')
    add_numbered_list(doc, [
        '个人信息与服务项目管理：服务人员可维护个人资料，设置可提供的服务类型、服务区域、可服务日期与时段，发布与管理个人服务项目（含单次卡/多次卡、分级定价）。',
        '订单接收与处理功能：服务人员可查看系统分配或用户指定的订单，并进行接单、拒单（需填写原因）以及服务状态更新等操作，支持组团游订单的接受/拒绝/退出。',
        '服务状态管理功能：服务人员可对服务过程进行管理，包括开始服务、暂停服务、恢复服务、结束服务以及查看当前服务进度，可设置工作/休息模式及临时休息时段。',
        '收入与服务记录查询功能：服务人员可查看历史服务订单、个人收入统计、仪表盘数据以及用户评价信息。',
        '用户评价与反馈功能：服务人员在服务完成后可对用户配合度、沟通表现等进行评价，形成双向评价机制，并可向平台提交问题反馈或异常情况说明。',
        '消息通知功能：服务人员可接收新订单、派单、组团邀请、评价提醒等通知消息。',
    ])

    add_heading2(doc, '1.2 系统功能模块设计')
    add_note(doc, '（功能图必须有，使用专用工具绘制）')
    add_body(doc, '系统功能模块如下：')
    add_module_list(doc, [
        '登录模块：系统对用户、服务人员以及管理员的身份信息进行验证，支持微信授权登录、手机号验证码登录及账号密码登录，在验证成功后根据用户角色跳转至对应的操作界面，实现不同权限功能的访问。',
        '服务类型与服务信息管理模块：管理员对平台中的陪诊、陪护、陪玩等服务类型及具体服务项目进行新增、修改、删除以及上下架管理，并对服务价格、服务等级定价、服务时长等信息进行维护。',
        '用户与服务人员管理模块：管理员对平台用户及服务人员信息进行录入、查询、修改以及状态管理，对违规账号可进行禁用或恢复操作，同时支持服务人员等级规则配置与自动升降级。',
        '订单信息管理模块：管理员对用户预约订单、服务状态、人工派单以及退款申请进行统一管理，可查看订单详情并对异常订单进行处理。',
        '服务预约与智能分配模块：用户可根据服务类型、预约时间、服务区域以及服务人员等级进行预约下单，系统支持用户自主选择服务人员或由系统根据等级、评分、空闲状态自动进行智能分配。',
        '服务处理模块：服务人员可查看系统分配或用户指定的订单，并进行接单、拒单、开始服务、暂停/恢复、结束服务等操作，系统同步更新订单状态并记录服务过程。',
        '评价与反馈模块：用户在服务结束后可对服务人员进行多维度评分与评价，服务人员也可对用户进行反馈，管理员可对投诉、售后与意见信息进行查看与处理。',
        '订单与退款模块：用户可查看历史订单、取消订单以及提交退款申请，系统根据订单状态与服务完成情况进行退款金额计算，管理员审核后完成退款处理并生成相关记录。',
        '积分与营销模块：系统支持消费积分、积分商城兑换、优惠券发放与使用，以及组团游、新手礼包、限时折扣等营销活动管理。',
        '个人信息管理模块：用户与服务人员可对个人资料、联系方式、服务时间、工作模式等信息进行修改与维护。',
        '消息通知模块：系统向用户与服务人员推送订单状态、活动公告、派单提醒等消息，支持已读标记与关联业务跳转。',
        '数据统计分析模块：系统对平台中的订单数量、热门服务、服务人员评价、用户消费情况、退款率以及营收数据等信息进行统计与分析，方便管理员了解平台运营情况。',
    ])
    add_body(doc, '如图1-1所示。')
    add_figure_placeholder(doc, '图1-1 陪伴服务系统功能模块图')
    doc.add_page_break()


def build_chapter2_section1(doc):
    add_heading1(doc, '第2章 系统设计')
    add_note(doc, '（核心功能流程图必须有，使用专用工具绘制）')
    add_heading2(doc, '2.1 系统整体业务流程图')
    add_body(doc, '系统整体业务流程图')
    add_body(doc, (
        '用户进入微信小程序后，首先通过系统登录模块完成身份验证（微信授权、手机号或账号密码），'
        '系统根据用户角色进入对应功能界面。普通用户登录后，可浏览平台中的陪诊、陪护、陪玩、陪吃、陪游、'
        '陪学、陪聊等服务信息，并根据服务类型、价格、服务等级以及服务人员评价等条件进行查询与筛选。'
        '用户选择所需服务后，可进入服务详情页查看服务介绍、价格、时长及可选服务人员，填写预约时间、服务地点'
        '以及定制需求后提交订单，系统将预约信息录入订单表（orders），订单初始状态为"待支付"。'
        '用户完成支付后，订单状态更新为"待接单"。'
    ))
    add_body(doc, (
        '当用户提交订单后，系统根据用户选择情况进行服务人员分配：若用户指定服务人员，则系统判断该服务人员'
        '当前工作模式、休息时段及已有订单排期是否冲突；若未指定服务人员，则系统根据派单规则，'
        '综合服务人员等级、平均评分、空闲状态、服务区域匹配度等条件进行智能匹配。'
        '订单生成并分配完成后，订单状态保持"待接单"，并向对应服务人员发送任务通知（notifications 表）。'
    ))
    add_body(doc, (
        '服务人员登录系统后，可查看平台分配或用户指定的订单信息，并对订单进行接单或拒单处理。'
        '服务人员接单成功后，系统修改订单状态为"待服务"；服务开始时更新为"服务中"，'
        '支持暂停与恢复操作；服务完成后，用户与服务人员双方确认完成，系统自动更新订单状态为"已完成"，'
        '并在 order_services 表中生成服务记录，同时触发积分发放与收入统计。'
    ))
    add_body(doc, (
        '用户在服务结束后，可对服务人员进行综合评分及专业度、态度、准时等多维度评价，'
        '评价信息录入 reviews 表，用于后续服务人员等级评定与平台推荐排序。'
        '同时，服务人员也可对用户配合度、沟通表现进行评价，形成双向评价机制。'
        '若用户对服务存在不满，可提交售后申请（after_sales_requests 表），管理员可查看并处理，'
        '处理结果反馈给用户。若用户因特殊原因需要取消订单或申请退款，可提交退款申请，'
        '系统根据订单状态以及服务完成情况进行退款金额计算，由管理员在后台进行审核与处理。'
    ))
    add_body(doc, (
        '管理员登录后台管理系统后，可对平台服务类型与服务项目进行新增、修改、删除以及上下架管理，'
        '维护服务价格与服务等级定价信息；对平台用户及服务人员信息进行管理，包括信息查询、状态修改'
        '以及违规账号处理；对于服务订单，管理员可查看订单详情，进行人工派单或对异常订单进行干预处理，'
        '并可根据服务人员完成订单数量、用户评分以及服务质量调整服务人员等级。'
        '系统同时设置消息通知与营销活动功能，管理员可发布优惠券、组团游、新手礼包、限时折扣等活动，'
        '系统还会对订单数量、热门服务、服务人员排行、用户消费情况以及平台营收等数据进行统计分析，'
        '为管理员优化平台运营提供数据支持。如图2-1所示。'
    ))
    add_figure_placeholder(doc, '图2-1 系统整体业务流程图')


def build_chapter2_section2(doc):
    add_heading2(doc, '2.2 系统各业务流程图')
    flows = [
        ('1. 用户注册与登录操作流程', (
            '用户进入微信小程序后，首先进入登录界面，可选择微信授权登录、手机号验证码登录或账号密码登录完成身份认证。'
            '系统在后端 users 表中验证用户信息，验证成功后签发 JWT 令牌，根据用户角色（普通用户 role=1、'
            '服务人员 role=2、管理员 role=3）进入对应的功能页面。若用户为首次微信登录，系统自动创建用户记录'
            '并初始化积分账户（user_points 表）；若为已注册用户，则直接读取用户相关数据并进入系统主页。'
            '服务人员需额外完成实名注册（姓名、身份证号、手机号验证），信息写入 service_providers 表，'
            '经管理员审核通过后方可接单。如图2-2所示。'
        ), '图2-2 登录流程图'),
        ('2. 服务浏览与查询操作流程', (
            '用户登录系统后，可进入首页浏览平台提供的服务类型图标及热门推荐服务。'
            '用户可进入服务列表页，根据服务类型、价格区间、服务等级等条件进行筛选，'
            '也可通过搜索页输入关键词检索服务项目或服务人员。'
            '系统后端通过 ServiceController、SearchController 接口查询 service_types、services、'
            'service_providers 等表，并将符合条件的结果以 JSON 格式返回小程序端展示。如图2-3所示。'
        ), '图2-3 服务浏览与查询流程图'),
        ('3. 服务预约与下单业务操作流程', (
            '用户选择所需服务后，进入服务详情页面查看服务介绍、服务价格、服务时长、可预约日期时段'
            '以及可选服务人员列表，确认服务后填写预约日期、预约时段、服务地点及定制需求（如有），'
            '并选择是否指定服务人员。用户提交订单后，系统在后端 OrderController.createOrder 中生成唯一订单编号，'
            '写入 orders 表，并根据是否使用优惠券计算实付金额。用户完成模拟支付后，订单状态由"待支付"更新为"待接单"。'
            '若用户自主选择服务人员，系统调用 providerAvailability 工具判断该人员档期是否空闲；'
            '若用户未指定，则系统调用智能派单算法按等级、评分、空闲状态自动分配。如图2-4所示。'
        ), '图2-4 服务预约与下单流程图'),
        ('4. 服务接单操作流程', (
            '服务人员登录系统后，在"我的订单"页面查看系统分配或用户指定的待接订单。'
            '服务人员确认订单后进行接单操作，系统将订单状态修改为"待服务"并向用户发送通知；'
            '若服务人员拒单，需填写拒单原因（reject_reason 字段），系统重新触发智能派单或等待管理员人工派单。'
            '服务开始时，服务人员更新订单状态为"服务中"；服务过程中可暂停/恢复；'
            '服务结束后，用户与服务人员双方确认完成，系统将订单状态更新为"已完成"，'
            '在 order_services 表写入服务记录，并触发积分奖励与收入统计。如图2-5所示。'
        ), '图2-5 服务接单流程图'),
        ('5. 订单管理与退款业务操作流程', (
            '用户可在"我的订单"页面查看历史订单及当前订单状态，按状态 Tab 筛选（全部、待支付、待服务、'
            '服务中、已完成等）。若用户因特殊原因无法继续服务，可在允许的时间窗口内提交取消订单或退款申请，'
            '系统根据 AdminRefundController 中配置的退款规则及订单当前状态、服务完成情况计算退款金额，'
            '生成退款申请并更新订单 refund_result 为 pending。管理员进入后台退款管理页，'
            '对退款申请进行审核：审核通过后更新订单状态为"已退费"并完成退款操作；'
            '若审核未通过，填写拒绝原因并反馈给用户。如图2-6所示。'
        ), '图2-6 订单管理与退款流程图'),
        ('6. 评价与售后业务操作流程', (
            '服务完成后，用户可对服务人员进行综合评分及专业度、服务态度、准时程度等多维度评价，'
            '支持文字、图片及行为标签，评价信息写入 reviews 表（reviewer_type=user），'
            '并作为服务人员 avg_rating 更新及等级评定的重要依据。'
            '若用户对服务过程存在不满，可提交售后申请（类型：退款/提前结束/纠纷），'
            '系统将申请内容写入 after_sales_requests 表，管理员查看后进行同意或拒绝处理并填写回复。'
            '服务人员在服务结束后，也可对用户配合度、沟通表现进行评价（reviewer_type=provider），'
            '形成双向评价机制。如图2-7所示。'
        ), '图2-7 评价与售后流程图'),
        ('7. 服务人员管理业务操作流程', (
            '管理员登录后台管理系统后，可对服务人员信息进行新增、修改、查询以及启用/禁用管理。'
            '系统根据 system_configs 表中 provider_level_rules 配置，'
            '综合服务人员完成订单数量（total_services）、平均评分（avg_rating）及差评次数'
            '自动计算并更新服务等级（level 字段，1-10 星）。'
            '管理员还可查看服务人员当前工作模式（work_mode）、空闲状态（available）、'
            '服务区域及历史服务记录，对违规服务人员进行禁用或限制处理。如图2-8所示。'
        ), '图2-8 服务人员管理流程图'),
        ('8. 服务项目管理业务操作流程', (
            '管理员可对平台中的服务类型（service_types）与具体服务项目（services）进行统一管理，'
            '包括新增服务、修改服务内容、调整基础价格与等级定价（level_prices）、设置服务时长、'
            '可服务日期时段以及上下架操作。服务人员也可在"我的服务"中发布个人服务项目，'
            '经系统校验后写入 services 表并建立 provider_services 关联。'
            '系统将修改后的服务信息实时展示给用户端。如图2-9所示。'
        ), '图2-9 服务项目管理流程图'),
        ('9. 数据统计分析业务操作流程', (
            '系统对用户订单、服务记录、评价信息、退款记录及营销数据等进行实时统计。'
            '管理员可在后台统计页面查看销售趋势、热门服务排行、服务人员评分排行、'
            '用户消费分布、退款率、等级分布以及营收汇总等统计结果，支持按时间范围筛选与数据导出。'
            '统计数据通过 AdminStatisticsController 聚合查询 orders、reviews、service_providers 等表获得，'
            '为平台运营优化与服务调整提供参考依据。如图2-10所示。'
        ), '图2-10 数据统计分析流程图'),
        ('10. 营销活动与消息通知业务操作流程', (
            '管理员可通过后台营销模块创建满减券、折扣券、组团游、新手礼包、限时折扣等活动，'
            '活动信息写入 coupons 表及相关扩展字段。用户可在首页查看优惠活动，'
            '领取新手礼包（user_newuser_gift_log 防重复领取）或在积分商城兑换优惠券。'
            '下单时系统自动匹配可用优惠券并计算抵扣金额。系统通过 notifications 表向用户与服务人员'
            '推送订单状态变更、活动通知、派单提醒等消息，用户进入消息中心可查看并标记已读。如图2-11所示。'
        ), '图2-11 营销活动与消息通知流程图'),
    ]
    for title, desc, caption in flows:
        add_body(doc, title)
        add_body(doc, desc)
        add_figure_placeholder(doc, caption)


def build_chapter2_section3(doc):
    add_heading2(doc, '2.3 数据库设计')
    add_body(doc, '陪伴服务系统的E-R图如图2-12所示。')
    add_note(doc, '（E-R图必须有，使用专用工具绘制）')
    add_figure_placeholder(doc, '图2-12 陪伴服务系统E-R图')
    add_body(doc, '陪伴服务系统具体表结构如下：')
    add_note(doc, '（基本表结构必须有，标明主键、外键）')

    tables = [
        ('1. 用户信息表：', '平台普通用户的基本信息按微信 openid、昵称、手机号、角色等数据项保存于数据库中。如表2-1所示。',
         '表2-1 用户信息表 users', [
            ['id', '用户编号', 'INT', '否', '主键，自增'],
            ['openid', '微信openid', 'VARCHAR(100)', '否', '唯一索引'],
            ['nickname', '昵称', 'VARCHAR(100)', '是', ''],
            ['avatar_url', '头像地址', 'VARCHAR(500)', '是', ''],
            ['phone', '手机号', 'VARCHAR(20)', '是', ''],
            ['password', '登录密码', 'VARCHAR(255)', '是', ''],
            ['real_name', '真实姓名', 'VARCHAR(50)', '是', ''],
            ['id_card', '身份证号', 'VARCHAR(20)', '是', ''],
            ['gender', '性别', 'TINYINT', '是', '0未知1男2女'],
            ['role', '角色', 'TINYINT', '是', '1用户2服务人员3管理员'],
            ['service_types', '服务类型偏好', 'TEXT', '是', 'JSON数组'],
            ['status', '账号状态', 'TINYINT', '是', '0禁用1正常'],
            ['points', '积分余额', 'INT', '是', ''],
            ['total_consumed', '累计消费', 'DECIMAL(10,2)', '是', ''],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('2. 服务人员信息表：', '平台服务人员的基本信息、等级及统计数据按 openid、姓名、等级、评分等数据项保存于数据库中。如表2-2所示。',
         '表2-2 服务人员信息表 service_providers', [
            ['id', '人员编号', 'INT', '否', '主键，自增'],
            ['openid', '微信openid', 'VARCHAR(100)', '否', '唯一索引'],
            ['nickname', '昵称', 'VARCHAR(100)', '是', ''],
            ['phone', '手机号', 'VARCHAR(20)', '否', ''],
            ['real_name', '真实姓名', 'VARCHAR(50)', '否', ''],
            ['id_card', '身份证号', 'VARCHAR(20)', '否', '唯一索引'],
            ['level', '服务等级', 'INT', '是', '1-10星'],
            ['total_services', '完成服务数', 'INT', '是', ''],
            ['avg_rating', '平均评分', 'DECIMAL(3,2)', '是', ''],
            ['status', '账号状态', 'TINYINT', '是', '0禁用1正常2审核中'],
            ['available', '空闲状态', 'TINYINT', '是', '0忙碌1空闲'],
            ['work_mode', '工作模式', 'TINYINT', '是', '0休息1工作'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('3. 服务类型表：', '平台服务分类信息按类型名称、图标、描述等数据项保存于数据库中。如表2-3所示。',
         '表2-3 服务类型表 service_types', [
            ['id', '类型编号', 'INT', '否', '主键，自增'],
            ['name', '类型名称', 'VARCHAR(50)', '否', '如陪诊/陪玩等'],
            ['icon', '图标', 'VARCHAR(200)', '是', ''],
            ['description', '类型描述', 'TEXT', '是', ''],
            ['sort_order', '排序', 'INT', '是', ''],
            ['status', '状态', 'TINYINT', '是', '0下架1上架'],
            ['level_price_ranges', '等级价格区间', 'TEXT', '是', 'JSON'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('4. 服务项目表：', '具体服务项目信息按名称、类型、价格、时长等数据项保存于数据库中。如表2-4所示。',
         '表2-4 服务项目表 services', [
            ['id', '项目编号', 'INT', '否', '主键，自增'],
            ['provider_id', '所属服务人员', 'INT', '是', '外键→service_providers.id'],
            ['type_id', '服务类型', 'INT', '否', '外键→service_types.id'],
            ['name', '项目名称', 'VARCHAR(100)', '否', ''],
            ['description', '项目描述', 'TEXT', '是', ''],
            ['duration', '服务时长(分钟)', 'INT', '否', ''],
            ['base_price', '基础价格', 'DECIMAL(10,2)', '否', ''],
            ['level_prices', '分级定价', 'TEXT', '是', '铜银金JSON'],
            ['card_type', '卡类型', 'TINYINT', '是', '1单次2多次'],
            ['card_count', '可用次数', 'INT', '是', ''],
            ['weekdays', '可服务日期', 'TEXT', '是', 'JSON数组'],
            ['time_slots', '可服务时段', 'TEXT', '是', 'JSON数组'],
            ['service_area', '服务区域', 'VARCHAR(500)', '是', ''],
            ['status', '状态', 'TINYINT', '是', '0下架1上架'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('5. 服务套餐表：', '多次卡套餐信息按服务项目、次数、折扣及价格等数据项保存于数据库中。如表2-5所示。',
         '表2-5 服务套餐表 service_packages', [
            ['id', '套餐编号', 'INT', '否', '主键，自增'],
            ['service_id', '所属服务', 'INT', '否', '外键→services.id'],
            ['name', '套餐名称', 'VARCHAR(100)', '否', ''],
            ['count', '服务次数', 'INT', '否', ''],
            ['discount', '折扣比例', 'DECIMAL(5,2)', '是', ''],
            ['price', '套餐价格', 'DECIMAL(10,2)', '否', ''],
            ['status', '状态', 'TINYINT', '是', '0下架1上架'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('6. 订单信息表：', '用户预约订单按订单号、用户、服务人员、服务项目、预约时间及金额等数据项保存于数据库中。如表2-6所示。',
         '表2-6 订单信息表 orders', [
            ['id', '订单编号', 'INT', '否', '主键，自增'],
            ['order_no', '订单号', 'VARCHAR(32)', '否', '唯一索引'],
            ['user_id', '下单用户', 'INT', '否', '外键→users.id'],
            ['provider_id', '服务人员', 'INT', '是', '外键→service_providers.id'],
            ['service_id', '服务项目', 'INT', '否', '外键→services.id'],
            ['package_id', '套餐', 'INT', '是', '外键→service_packages.id'],
            ['scheduled_date', '预约日期', 'DATE', '否', ''],
            ['scheduled_time', '预约时段', 'VARCHAR(20)', '否', ''],
            ['total_price', '订单总额', 'DECIMAL(10,2)', '否', ''],
            ['paid_amount', '实付金额', 'DECIMAL(10,2)', '是', ''],
            ['discount_amount', '优惠券抵扣', 'DECIMAL(10,2)', '是', ''],
            ['status', '订单状态', 'TINYINT', '是', '0待支付~7已退费'],
            ['assign_type', '派单方式', 'TINYINT', '是', '0系统1自选'],
            ['refund_amount', '退款金额', 'DECIMAL(10,2)', '是', ''],
            ['refund_result', '退款结果', 'VARCHAR(20)', '是', ''],
            ['user_complete_confirmed', '用户确认完成', 'TINYINT', '是', ''],
            ['provider_complete_confirmed', '人员确认完成', 'TINYINT', '是', ''],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('7. 订单服务记录表：', '多次卡订单的每次服务使用记录按订单、服务日期、状态等数据项保存于数据库中。如表2-7所示。',
         '表2-7 订单服务记录表 order_services', [
            ['id', '记录编号', 'INT', '否', '主键，自增'],
            ['order_id', '所属订单', 'INT', '否', '外键→orders.id'],
            ['service_date', '服务日期', 'DATE', '是', ''],
            ['start_time', '开始时间', 'VARCHAR(20)', '是', ''],
            ['end_time', '结束时间', 'VARCHAR(20)', '是', ''],
            ['status', '使用状态', 'TINYINT', '是', '0未用1已用2取消'],
            ['used_by_provider', '实际服务人员', 'INT', '是', '外键→service_providers.id'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
        ]),
        ('8. 评价信息表：', '用户与服务人员的双向评价按订单、评分维度、内容等数据项保存于数据库中。如表2-8所示。',
         '表2-8 评价信息表 reviews', [
            ['id', '评价编号', 'INT', '否', '主键，自增'],
            ['order_id', '关联订单', 'INT', '否', '外键→orders.id'],
            ['reviewer_type', '评价方类型', 'VARCHAR(20)', '否', 'user/provider'],
            ['user_id', '订单用户', 'INT', '否', '外键→users.id'],
            ['provider_id', '服务人员', 'INT', '否', '外键→service_providers.id'],
            ['overall_rating', '综合评分', 'DECIMAL(3,1)', '是', ''],
            ['professional_rating', '专业度', 'DECIMAL(3,1)', '是', '用户评价'],
            ['attitude_rating', '服务态度', 'DECIMAL(3,1)', '是', ''],
            ['punctual_rating', '准时程度', 'DECIMAL(3,1)', '是', '用户评价'],
            ['cooperation_rating', '配合程度', 'DECIMAL(3,1)', '是', '人员评价'],
            ['communication_rating', '沟通表现', 'DECIMAL(3,1)', '是', '人员评价'],
            ['content', '评价内容', 'TEXT', '是', ''],
            ['images', '评价图片', 'TEXT', '是', 'JSON数组'],
            ['status', '显示状态', 'TINYINT', '是', '0隐藏1显示'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
        ]),
        ('9. 优惠券表：', '平台优惠券及营销活动信息按名称、类型、面值、有效期等数据项保存于数据库中。如表2-9所示。',
         '表2-9 优惠券表 coupons', [
            ['id', '优惠券编号', 'INT', '否', '主键，自增'],
            ['name', '优惠券名称', 'VARCHAR(100)', '否', ''],
            ['type', '券类型', 'TINYINT', '是', '1满减2折扣3组团等'],
            ['discount_value', '优惠值', 'DECIMAL(10,2)', '否', ''],
            ['min_amount', '最低消费', 'DECIMAL(10,2)', '是', ''],
            ['valid_start', '生效日期', 'DATE', '否', ''],
            ['valid_end', '失效日期', 'DATE', '否', ''],
            ['total_count', '发放总量', 'INT', '是', ''],
            ['used_count', '已使用数', 'INT', '是', ''],
            ['points_cost', '兑换积分', 'INT', '是', '积分商城'],
            ['status', '状态', 'TINYINT', '是', '0停用1启用'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('10. 用户优惠券表：', '用户领取的优惠券记录按用户、券、使用状态等数据项保存于数据库中。如表2-10所示。',
         '表2-10 用户优惠券表 user_coupons', [
            ['id', '记录编号', 'INT', '否', '主键，自增'],
            ['user_id', '用户', 'INT', '否', '外键→users.id'],
            ['coupon_id', '优惠券', 'INT', '是', '外键→coupons.id'],
            ['status', '使用状态', 'TINYINT', '是', '1未用2已用3过期'],
            ['used_order_id', '使用订单', 'INT', '是', '外键→orders.id'],
            ['expire_time', '过期时间', 'DATETIME', '是', ''],
            ['created_at', '领取时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('11. 用户积分表：', '用户积分账户按当前积分、累计获得与消耗等数据项保存于数据库中。如表2-11所示。',
         '表2-11 用户积分表 user_points', [
            ['id', '记录编号', 'INT', '否', '主键，自增'],
            ['user_id', '用户', 'INT', '否', '外键→users.id，唯一'],
            ['points', '当前积分', 'INT', '是', ''],
            ['total_earned', '累计获得', 'INT', '是', ''],
            ['total_spent', '累计消耗', 'INT', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('12. 积分变动记录表：', '积分获得与消耗明细按用户、类型、数量及来源等数据项保存于数据库中。如表2-12所示。',
         '表2-12 积分变动记录表 point_transactions', [
            ['id', '记录编号', 'INT', '否', '主键，自增'],
            ['user_id', '用户', 'INT', '否', '外键→users.id'],
            ['type', '变动类型', 'TINYINT', '否', '1获得2消耗'],
            ['amount', '变动数量', 'INT', '否', ''],
            ['source', '来源描述', 'VARCHAR(100)', '是', ''],
            ['order_id', '关联订单', 'INT', '是', '外键→orders.id'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
        ]),
        ('13. 管理员表：', '后台管理员账号按用户名、密码、昵称等数据项保存于数据库中。如表2-13所示。',
         '表2-13 管理员表 admins', [
            ['id', '管理员编号', 'INT', '否', '主键，自增'],
            ['username', '用户名', 'VARCHAR(50)', '否', '唯一索引'],
            ['password', '密码', 'VARCHAR(255)', '否', ''],
            ['nickname', '昵称', 'VARCHAR(50)', '是', ''],
            ['status', '状态', 'TINYINT', '是', '0禁用1正常'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('14. 系统通知表：', '向用户或服务人员的通知消息按标题、内容、类型及已读状态等数据项保存于数据库中。如表2-14所示。',
         '表2-14 系统通知表 notifications', [
            ['id', '通知编号', 'INT', '否', '主键，自增'],
            ['user_id', '目标用户', 'INT', '是', '外键→users.id'],
            ['provider_id', '目标服务人员', 'INT', '是', '外键→service_providers.id'],
            ['title', '通知标题', 'VARCHAR(100)', '否', ''],
            ['content', '通知内容', 'TEXT', '否', ''],
            ['type', '通知类型', 'TINYINT', '是', '1系统2订单3评价'],
            ['read', '已读状态', 'TINYINT', '是', '0未读1已读'],
            ['ref_type', '关联类型', 'VARCHAR(50)', '是', ''],
            ['ref_id', '关联ID', 'INT', '是', ''],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
        ]),
        ('15. 售后申请表：', '用户售后与投诉申请按订单、类型、原因及处理结果等数据项保存于数据库中。如表2-15所示。',
         '表2-15 售后申请表 after_sales_requests', [
            ['id', '申请编号', 'INT', '否', '主键，自增'],
            ['order_id', '关联订单', 'INT', '否', '外键→orders.id'],
            ['user_id', '申请用户', 'INT', '否', '外键→users.id'],
            ['type', '申请类型', 'VARCHAR(30)', '是', 'refund/end_early/dispute'],
            ['reason', '申请原因', 'TEXT', '否', ''],
            ['images', '凭证图片', 'TEXT', '是', 'JSON数组'],
            ['status', '处理状态', 'TINYINT', '是', '0待处理1同意2拒绝'],
            ['admin_reply', '管理员回复', 'TEXT', '是', ''],
            ['refund_amount', '退款金额', 'DECIMAL(10,2)', '是', ''],
            ['created_at', '申请时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('16. 服务人员服务项目关联表：', '服务人员与服务项目的多对多关联关系保存于数据库中。如表2-16所示。',
         '表2-16 服务人员服务项目关联表 provider_services', [
            ['id', '记录编号', 'INT', '否', '主键，自增'],
            ['provider_id', '服务人员', 'INT', '否', '外键→service_providers.id'],
            ['service_id', '服务项目', 'INT', '否', '外键→services.id'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
        ]),
        ('17. 服务人员日程表：', '服务人员每日可服务时段按人员、日期及时段等数据项保存于数据库中。如表2-17所示。',
         '表2-17 服务人员日程表 provider_schedule', [
            ['id', '记录编号', 'INT', '否', '主键，自增'],
            ['provider_id', '服务人员', 'INT', '否', '外键→service_providers.id'],
            ['date', '日期', 'DATE', '否', ''],
            ['time_slots', '可用时段', 'TEXT', '是', 'JSON数组'],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
        ('18. 系统配置表：', '平台全局配置参数按配置键、配置值及描述等数据项保存于数据库中。如表2-18所示。',
         '表2-18 系统配置表 system_configs', [
            ['id', '配置编号', 'INT', '否', '主键，自增'],
            ['config_key', '配置键', 'VARCHAR(100)', '否', '唯一索引'],
            ['config_value', '配置值', 'TEXT', '是', 'JSON格式'],
            ['description', '配置描述', 'VARCHAR(500)', '是', ''],
            ['created_at', '创建时间', 'DATETIME', '是', ''],
            ['updated_at', '更新时间', 'DATETIME', '是', ''],
        ]),
    ]

    for prefix, desc, title, cols in tables:
        add_body(doc, prefix)
        add_table_def(doc, title, desc, cols)

    doc.add_page_break()


def build_chapter3(doc):
    add_heading1(doc, '第3章 系统实现')
    add_heading2(doc, '3.1 系统开发环境')
    add_note(doc, '（根据实际情况写）')
    env_items = [
        ('开发用计算机系统', 'Windows 10 / Windows 11'),
        ('后端运行环境', 'Node.js 18+'),
        ('后端开发框架', 'Express 4.x'),
        ('前端开发平台', '微信开发者工具（微信小程序）'),
        ('编程语言', 'JavaScript（ES6+）、WXML、WXSS、SQL'),
        ('数据库', 'MySQL 8.0'),
        ('数据库管理工具', 'MySQL Workbench / Navicat'),
        ('接口测试工具', 'Postman / 微信开发者工具 Network 面板'),
        ('版本管理', 'Git'),
        ('项目启动方式', '运行 start.bat 一键启动 MySQL 与后端服务，微信开发者工具打开项目根目录编译小程序'),
    ]
    for k, v in env_items:
        add_body(doc, f'{k}：{v}')

    add_heading2(doc, '3.2 系统功能实现')
    add_note(doc, '（功能截图必须有，并用文字说明）')

    impls = [
        ('1. 用户登录界面：', (
            '用户进入小程序后首先进入登录页面（pages/login-user），可选择微信授权登录、手机号验证码登录或账号密码登录。'
            '前端将登录请求发送至后端 /api/user/login、/api/user/login-by-phone 或 /api/user/login-by-password 接口，'
            '后端 UserController 在 users 表中查询并验证用户信息，验证成功后使用 jsonwebtoken 签发 JWT 令牌返回前端，'
            '前端将 token 存入本地存储并在后续请求的 Authorization 头中携带。'
            '服务人员登录（pages/login-provider）与管理员登录（pages/login-admin）采用类似机制，'
            '分别查询 service_providers 表与 admins 表，验证成功后跳转至对应功能主页。如图3-1所示。'
        ), '图3-1 用户登录页面'),
        ('2. 首页与服务浏览页面：', (
            '用户登录后进入首页（pages/index），页面通过 /api/services/types 接口加载服务类型图标列表，'
            '通过 /api/services/hot 与 /api/services/recommend 接口加载热门推荐服务。'
            '用户点击服务类型进入服务列表页（pages/service-list），可按类型筛选；'
            '进入搜索页（pages/search）可输入关键词检索服务项目或服务人员，'
            '后端 SearchController 对 services、service_providers 表进行模糊查询并返回结果。如图3-2所示。'
        ), '图3-2 首页与服务浏览页面'),
        ('3. 服务详情与预约下单页面：', (
            '用户点击服务项目进入详情页（pages/service-detail），通过 /api/services/:id 接口获取服务名称、'
            '描述、价格、时长、可预约日期时段及可选服务人员列表。用户确认后填写预约日期、时段、服务地点，'
            '可选择指定服务人员或系统自动分配，点击下单调用 /api/orders 接口，'
            '后端 OrderController.createOrder 生成唯一 order_no 写入 orders 表。'
            '用户进入订单详情页（pages/order-detail）完成模拟支付（/api/orders/:id/pay），'
            '订单状态由待支付更新为待接单，系统向服务人员发送通知。如图3-3所示。'
        ), '图3-3 服务详情与预约下单页面'),
        ('4. 订单管理页面：', (
            '用户在"我的订单"页（pages/orders）通过 /api/orders 接口按状态 Tab 查看全部、待支付、待服务、'
            '服务中、已完成等订单列表。点击订单进入详情页可查看订单编号、服务项目、预约时间、服务人员、'
            '金额及当前状态，并执行取消订单（/api/orders/:id/cancel）、申请退款（/api/orders/refund）、'
            '确认完成（/api/orders/:id/confirm-complete）等操作。多次卡订单还支持预约下次服务'
            '（/api/orders/:id/book-next）。如图3-4所示。'
        ), '图3-4 订单管理页面'),
        ('5. 服务评价页面：', (
            '订单完成后，用户可在评价页（pages/review）对服务人员进行综合评分及专业度、服务态度、准时程度等'
            '多维度打分，填写文字评价并上传图片。前端调用 /api/reviews 接口，'
            '后端 ReviewController 将评价写入 reviews 表（reviewer_type=user），'
            '并更新 service_providers 表的 avg_rating 字段。用户还可在"我的评价"页（pages/my-reviews）'
            '查看历史评价记录。如图3-5所示。'
        ), '图3-5 服务评价页面'),
        ('6. 积分与优惠券页面：', (
            '用户在积分页（pages/points）通过 /api/user/points 查看当前积分余额，'
            '在积分明细页（pages/points-record）查看 point_transactions 表中的变动记录。'
            '积分商城页（pages/points-mall）展示可兑换商品与优惠券，用户调用 /api/user/exchange-coupon 兑换后'
            '写入 user_coupons 表。优惠券页（pages/coupons）展示已领取的优惠券及使用状态，'
            '下单时可调用 /api/orders/checkout-coupons 获取可用券列表并抵扣金额。如图3-6所示。'
        ), '图3-6 积分与优惠券页面'),
        ('7. 服务人员工作台页面：', (
            '服务人员登录后进入工作台（pages/provider-home），通过 /api/provider/dashboard 接口获取待接订单数、'
            '今日收入、平均评分等概览数据。在订单页（pages/provider-orders）查看待接/进行中/已完成订单，'
            '可执行接单、拒单、开始服务、暂停/恢复、结束服务等操作，'
            '每次状态变更调用 /api/provider/order/status 接口更新 orders 表状态并发送用户通知。如图3-7所示。'
        ), '图3-7 服务人员工作台页面'),
        ('8. 服务人员服务项目管理页面：', (
            '服务人员在"我的服务"页（pages/provider-my-services）查看已发布的服务项目列表，'
            '点击编辑进入 provider-service-edit 页面，可设置服务名称、类型、价格、分级定价、'
            '可服务日期时段、服务区域及单次卡/多次卡类型。保存时调用 /api/provider/my-services 接口，'
            '后端将数据写入 services 表并维护 provider_services 关联关系。如图3-8所示。'
        ), '图3-8 服务人员服务项目管理页面'),
        ('9. 管理员后台首页：', (
            '管理员通过 login-admin 页面登录后进入管理后台（pages/admin），'
            '调用 /api/admin/stats 接口展示平台订单总数、用户总数、服务人员总数、今日营收等统计卡片，'
            '并提供服务管理、人员管理、订单管理、退款审核、营销管理、数据统计等功能入口。如图3-9所示。'
        ), '图3-9 管理员后台首页'),
        ('10. 管理员服务类型与服务项目管理页面：', (
            '管理员在服务管理页（pages/admin-services）可查看全部服务项目，'
            '进入服务类型管理（pages/admin-services/types）对陪诊、陪玩等类型进行新增、编辑、上下架。'
            '进入项目详情页可修改服务名称、价格、时长、等级定价及状态，'
            '调用 /api/admin/services 与 /api/admin/service-types 系列接口操作数据库。如图3-10所示。'
        ), '图3-10 管理员服务管理页面'),
        ('11. 管理员订单管理与人工派单页面：', (
            '管理员在订单管理页（pages/admin-orders）查看全部订单，可按状态、日期筛选。'
            '进入订单详情页可查看完整订单信息，对未分配服务人员的订单进入派单页（pages/admin-orders/assign），'
            '从可用服务人员列表中选择并调用 /api/admin/orders/:id/assign 接口完成人工派单。'
            '对异常订单可强制取消或修改状态。如图3-11所示。'
        ), '图3-11 管理员订单管理与派单页面'),
        ('12. 管理员退款与售后审核页面：', (
            '管理员在退款管理页（pages/admin-refunds）查看用户提交的退款申请列表，'
            '进入详情页审核：同意则调用 /api/admin/refunds/:id/approve 更新订单状态并记录退款金额；'
            '拒绝则调用 reject 接口并填写拒绝原因。售后管理页（pages/admin-after-sales）'
            '处理用户提交的提前结束、纠纷等售后申请，调用 /api/admin/after-sales/:id/approve 或 reject 接口。'
            '退款规则可在 rules 页面配置并写入 system_configs 表。如图3-12所示。'
        ), '图3-12 管理员退款与售后审核页面'),
        ('13. 管理员营销管理与数据统计页面：', (
            '管理员在营销管理页（pages/admin-marketing）可创建和管理满减券、折扣券、组团游、'
            '新手礼包、限时折扣等活动，调用 /api/admin/marketing 系列接口写入 coupons 表。'
            '数据统计页（pages/admin-statistics）展示销售趋势折线图、热门服务排行、服务人员评分排行、'
            '用户消费分布、退款率及等级分布等图表，数据来源于 AdminStatisticsController 对 orders、'
            'reviews、service_providers 等表的聚合查询，支持导出。如图3-13所示。'
        ), '图3-13 管理员营销管理与数据统计页面'),
    ]

    for title, desc, caption in impls:
        add_body(doc, title)
        add_body(doc, desc)
        add_figure_placeholder(doc, caption)


def verify_no_library_terms(docx_path):
    """校验报告中不含图书管理系统残留用语"""
    doc = Document(docx_path)
    text = '\n'.join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += '\n' + cell.text
    forbidden = [
        '图书', '借阅', '读者', '书籍', '图书馆', 'BookInfo', 'ReaderInfo',
        'BorrowInfo', 'ReturnInfo', 'Metrogrid', 'Visual Studio', 'SQL Server',
        'C#', '窗体', 'Datatable', '罚金', '借出', '归还', 'GetRandomCode',
        'CreateImage', 'Timer控件', '图书信息管理系统', '图书管理系统',
    ]
    found = [term for term in forbidden if term in text]
    if found:
        raise RuntimeError(f'报告仍含图书系统残留: {", ".join(found)}')
    print('内容校验通过：无图书管理系统残留')


def sync_root_reference_doc(docx_path):
    """用完整陪伴服务说明书覆盖项目根目录原参考 doc（内含图书系统章节）"""
    try:
        import win32com.client
    except ImportError:
        print('未安装 pywin32，跳过根目录 .doc 同步（请手动用 docx 覆盖参考 doc）')
        return

    root = os.path.dirname(OUTPUT_DIR)
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(docx_path)
    synced = []
    for name in os.listdir(root):
        if name.endswith('.doc') and not name.startswith('~$'):
            out = os.path.join(root, name)
            doc.SaveAs2(out, FileFormat=0)
            synced.append(out)
    companion_doc = os.path.join(root, '陪伴服务系统说明书.doc')
    if companion_doc not in synced:
        doc.SaveAs2(companion_doc, FileFormat=0)
        synced.append(companion_doc)
    doc.Close(False)
    word.Quit()
    for path in synced:
        print(f'已同步覆盖: {path}')


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    build_cover(doc)
    build_toc(doc)
    build_chapter1(doc)
    build_chapter2_section1(doc)
    build_chapter2_section2(doc)
    build_chapter2_section3(doc)
    build_chapter3(doc)

    doc.save(OUTPUT_FILE)
    print(f'报告已生成: {OUTPUT_FILE}')
    verify_no_library_terms(OUTPUT_FILE)
    sync_root_reference_doc(OUTPUT_FILE)


if __name__ == '__main__':
    main()
